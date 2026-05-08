"""
Generate 3 sample resumes (DOCX) to test the AI Resume Matcher app against the
AI/ML Engineer – Resume Intelligence Platform JD.

Expected outcomes:
  1. arjun_sharma_STRONG.docx   → Strong Fit  (≥75)  → Shortlist
  2. priya_mehta_AVERAGE.docx   → Moderate Fit (42-60) → Manual Review
  3. rohan_verma_LOW.docx       → Low Fit      (<42)  → Reject
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_bold_label(doc, label, content):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(content)
    return p


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


# ─────────────────────────────────────────────────────────────────────────────
# RESUME 1 – STRONG FIT (Shortlist)
# Hits ALL required skills + most preferred + good-to-haves
# 2 years relevant experience, quantified achievements, correct domain
# ─────────────────────────────────────────────────────────────────────────────
def create_strong_fit(path):
    doc = Document()

    # Header
    doc.add_heading("Arjun Sharma", 0)
    doc.add_paragraph("arjun.sharma@email.com | +91-9876543210 | Bengaluru, India")
    doc.add_paragraph("github.com/arjunsharma | linkedin.com/in/arjunsharma")

    # Summary
    add_heading(doc, "Professional Summary", 2)
    doc.add_paragraph(
        "AI/ML Engineer with 2.5 years of experience building NLP-powered resume screening, "
        "job matching, and candidate evaluation systems. Proficient in Python, FastAPI, and "
        "machine learning pipelines. Experienced with Generative AI and LLM integration "
        "(Gemini API) for explainable AI features. Deployed full-stack ML applications on "
        "Render and Vercel with React + Tailwind CSS frontends."
    )

    # Skills
    add_heading(doc, "Technical Skills", 2)
    add_bold_label(doc, "Languages & Frameworks", "Python, FastAPI, React, Tailwind CSS, Tailwind")
    add_bold_label(doc, "NLP & ML", "NLP, Natural Language Processing, Machine Learning, Text Preprocessing, Semantic Similarity, TF-IDF, TF IDF, scikit-learn, spaCy, Text Classification, Prompt Engineering")
    add_bold_label(doc, "APIs & Data", "REST API, REST APIs, JSON, JSON data handling, Git, Resume Parsing, Skill Extraction")
    add_bold_label(doc, "Generative AI", "Generative AI, Gemini, Gemini API, LLM, LLM Integration, Large Language Models")
    add_bold_label(doc, "Deployment", "Render, Vercel, Docker, Deployment")
    add_bold_label(doc, "Other", "Scoring Systems, Scoring System, ATS, Explainable AI, Evaluation Metrics, Evidence Validation, Gap Detection, Cosine Similarity")

    # Experience
    add_heading(doc, "Work Experience", 2)
    doc.add_paragraph("AI/ML Engineer – HireIQ Technologies, Bengaluru (Jan 2024 – Present)")
    add_bullet(doc, "Built resume-to-job matching pipeline using TF-IDF and semantic similarity, achieving 87% precision on internal benchmark.")
    add_bullet(doc, "Designed NLP-based skill extraction system using spaCy and custom regex patterns; reduced false-positive rate by 35%.")
    add_bullet(doc, "Developed FastAPI backend with 12+ REST API endpoints for ML-driven scoring, gap detection, and evidence validation.")
    add_bullet(doc, "Integrated Gemini LLM for generating explainable recruiter-facing summaries and candidate feedback.")
    add_bullet(doc, "Implemented prompt engineering workflows for resume tailoring and ATS optimization features.")
    add_bullet(doc, "Deployed full-stack application on Render (backend) and Vercel (React + Tailwind CSS frontend).")

    doc.add_paragraph("ML Engineer Intern – RecruitAI Labs, Remote (Jul 2023 – Dec 2023)")
    add_bullet(doc, "Built text classification model to categorize resume sections with 91% accuracy.")
    add_bullet(doc, "Worked with JSON/data handling pipelines to process 50,000+ resumes for scoring system evaluation.")
    add_bullet(doc, "Created scoring system for candidate fit analysis integrating required vs. preferred skill weighting.")
    add_bullet(doc, "Contributed to ATS simulation module to predict shortlist probability for candidates.")

    # Projects
    add_heading(doc, "Projects", 2)
    doc.add_paragraph("AI Resume Matcher (End-to-End ML Application)")
    add_bullet(doc, "Designed and deployed a full-stack AI resume matching platform with resume parsing, skill extraction, semantic similarity scoring, and LLM-based tailoring.")
    add_bullet(doc, "Stack: Python, FastAPI, React, Tailwind CSS, Gemini API, deployed on Render + Vercel.")
    add_bullet(doc, "Achieved end-to-end match analysis in under 3 seconds with explainable AI outputs.")

    doc.add_paragraph("NLP Pipeline for Unstructured Resume Data")
    add_bullet(doc, "Built text preprocessing and section parsing pipeline using semantic similarity and machine learning classifiers.")
    add_bullet(doc, "Handled structured and unstructured data with JSON serialization and Git-based version control.")

    # Education
    add_heading(doc, "Education", 2)
    doc.add_paragraph("B.Tech in Computer Science – IIT Hyderabad (2019 – 2023) | CGPA: 8.7/10")

    # Certifications
    add_heading(doc, "Certifications", 2)
    add_bullet(doc, "Google Professional Machine Learning Engineer Certification (2024)")
    add_bullet(doc, "DeepLearning.AI – NLP Specialization (Coursera, 2023)")
    add_bullet(doc, "FastAPI – Full Course for Beginners (Udemy, 2023)")

    doc.save(path)
    print(f"Created: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# RESUME 2 – AVERAGE FIT (Manual Review)
# Hits ~50-60% required skills, minimal preferred, adjacent domain
# Some Python/ML but no FastAPI, no NLP depth, no LLM/Generative AI
# ─────────────────────────────────────────────────────────────────────────────
def create_average_fit(path):
    doc = Document()

    # Header
    doc.add_heading("Priya Mehta", 0)
    doc.add_paragraph("priya.mehta@email.com | +91-9845671230 | Pune, India")
    doc.add_paragraph("github.com/priyamehta | linkedin.com/in/priyamehta")

    # Summary
    add_heading(doc, "Professional Summary", 2)
    doc.add_paragraph(
        "Python developer with 2 years of experience in machine learning, NLP, and "
        "REST API development. Built text preprocessing and text classification pipelines, "
        "applied semantic similarity for document matching, and developed FastAPI-based "
        "scoring systems for ranking outputs. Familiar with ATS workflows and basic "
        "resume parsing concepts. Seeking to deepen experience with LLM integration and "
        "Generative AI systems."
    )

    # Skills
    add_heading(doc, "Technical Skills", 2)
    add_bold_label(doc, "Languages", "Python, SQL, JavaScript (basic)")
    add_bold_label(doc, "ML & NLP", "Machine Learning, NLP, Natural Language Processing, Text Preprocessing, Text Classification, Semantic Similarity, scikit-learn, pandas, numpy")
    add_bold_label(doc, "APIs & Tools", "FastAPI, REST API, REST APIs, Git, JSON, JSON data handling, Flask")
    add_bold_label(doc, "Systems", "Scoring Systems, ATS, Resume Parsing (basic), Skill Extraction (basic)")
    add_bold_label(doc, "Web", "React (basic), HTML, CSS")

    # Experience
    add_heading(doc, "Work Experience", 2)
    doc.add_paragraph("ML Developer – DataSmart Solutions, Pune (Aug 2024 – Present)")
    add_bullet(doc, "Built NLP pipeline for text preprocessing and text classification of customer support tickets using Python and scikit-learn; improved routing accuracy by 31%.")
    add_bullet(doc, "Designed scoring system to rank candidate responses using semantic similarity (TF-IDF cosine similarity); reduced manual review time by 25%.")
    add_bullet(doc, "Developed FastAPI REST API to serve ML model predictions and scoring system outputs; handled JSON data handling throughout.")
    add_bullet(doc, "Built resume parsing module to extract structured information from unstructured text using Python and regex NLP techniques.")
    add_bullet(doc, "Used Git for version control; maintained JSON-based data pipelines for training and evaluation.")
    add_bullet(doc, "Integrated basic ATS-aware keyword extraction to flag resume-JD mismatches for the HR team.")

    doc.add_paragraph("Data Analyst Intern – Analytics Hub, Remote (Jan 2024 – Jul 2024)")
    add_bullet(doc, "Built text classification model for customer complaint categorization using NLP and machine learning; achieved 79% accuracy.")
    add_bullet(doc, "Performed text preprocessing (tokenization, stop-word removal, lemmatization) on 20,000+ records.")
    add_bullet(doc, "Consumed REST APIs and managed JSON data pipelines for automated reporting scripts.")
    add_bullet(doc, "Maintained Git repository for 3-person ML team.")

    # Projects
    add_heading(doc, "Projects", 2)
    doc.add_paragraph("Job-Resume Matching Tool (NLP)")
    add_bullet(doc, "Built a Python-based NLP tool to extract skills from resumes and compute semantic similarity scores against job descriptions using TF-IDF.")
    add_bullet(doc, "Designed scoring system to rank resumes; deployed via FastAPI REST API with JSON response format.")
    add_bullet(doc, "Used Git for version control; basic resume parsing and skill extraction logic implemented.")

    doc.add_paragraph("Text Classification Pipeline")
    add_bullet(doc, "Built machine learning text classification pipeline using Python, scikit-learn, text preprocessing; achieved 82% accuracy on held-out test set.")

    # Education
    add_heading(doc, "Education", 2)
    doc.add_paragraph("B.Sc in Computer Science – Pune University (2020 – 2023) | CGPA: 7.4/10")

    # Certifications
    add_heading(doc, "Certifications", 2)
    add_bullet(doc, "Natural Language Processing with Python – Coursera (2024)")
    add_bullet(doc, "IBM Machine Learning Professional Certificate (2023)")

    doc.save(path)
    print(f"Created: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# RESUME 3 – LOW FIT (Reject)
# Completely different domain – HR/Recruitment consultant background
# No Python, no ML, no NLP, no FastAPI, no relevant tech
# ─────────────────────────────────────────────────────────────────────────────
def create_low_fit(path):
    doc = Document()

    # Header
    doc.add_heading("Rohan Verma", 0)
    doc.add_paragraph("rohan.verma@email.com | +91-9712345678 | Mumbai, India")
    doc.add_paragraph("linkedin.com/in/rohanverma")

    # Summary
    add_heading(doc, "Professional Summary", 2)
    doc.add_paragraph(
        "HR Recruiter with 3 years of experience in talent acquisition, candidate screening, "
        "and HR operations across IT and non-IT sectors. Skilled in using ATS platforms, "
        "conducting interviews, and managing end-to-end recruitment cycles. Strong "
        "communication and stakeholder management skills."
    )

    # Skills
    add_heading(doc, "Technical Skills", 2)
    add_bold_label(doc, "Recruitment", "Talent Acquisition, Bulk Hiring, Campus Recruitment, Headhunting")
    add_bold_label(doc, "ATS Tools", "Naukri, LinkedIn Recruiter, Greenhouse ATS, Workday")
    add_bold_label(doc, "HR Tools", "SAP HR, MS Excel, Google Workspace")
    add_bold_label(doc, "Soft Skills", "Communication, Negotiation, Stakeholder Management, Team Coordination")

    # Experience
    add_heading(doc, "Work Experience", 2)
    doc.add_paragraph("Senior HR Recruiter – GlobalTech Staffing, Mumbai (Mar 2023 – Present)")
    add_bullet(doc, "Managed end-to-end recruitment for 50+ positions across IT and BFSI sectors.")
    add_bullet(doc, "Screened 500+ resumes monthly using Greenhouse ATS and LinkedIn Recruiter.")
    add_bullet(doc, "Coordinated with hiring managers to define job requirements and draft JDs.")
    add_bullet(doc, "Reduced time-to-hire by 20% by optimizing sourcing channels.")
    add_bullet(doc, "Conducted structured interviews and assessed cultural fit for senior roles.")

    doc.add_paragraph("HR Recruiter – QuickHire Consultants, Mumbai (Jun 2021 – Feb 2023)")
    add_bullet(doc, "Sourced candidates via Naukri, LinkedIn, and job fairs for IT and finance roles.")
    add_bullet(doc, "Managed offer rollouts, salary negotiations, and background verification processes.")
    add_bullet(doc, "Maintained candidate pipelines and updated recruitment trackers in MS Excel.")

    # Projects / Initiatives
    add_heading(doc, "HR Initiatives", 2)
    doc.add_paragraph("Campus Recruitment Drive – IIT Bombay (2024)")
    add_bullet(doc, "Led campus hiring drive and onboarded 15 engineering graduates.")

    doc.add_paragraph("ATS Implementation Project")
    add_bullet(doc, "Assisted in configuring Greenhouse ATS workflows for the recruitment team.")

    # Education
    add_heading(doc, "Education", 2)
    doc.add_paragraph("MBA in Human Resource Management – Symbiosis Institute, Pune (2018 – 2020) | CGPA: 7.1/10")
    doc.add_paragraph("B.Com – Mumbai University (2015 – 2018) | CGPA: 6.8/10")

    # Certifications
    add_heading(doc, "Certifications", 2)
    add_bullet(doc, "SHRM Certified Professional (SHRM-CP) – 2022")
    add_bullet(doc, "HR Analytics – LinkedIn Learning (2023)")

    doc.save(path)
    print(f"Created: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    create_strong_fit(os.path.join(OUTPUT_DIR, "arjun_sharma_STRONG_FIT.docx"))
    create_average_fit(os.path.join(OUTPUT_DIR, "priya_mehta_AVERAGE_FIT.docx"))
    create_low_fit(os.path.join(OUTPUT_DIR, "rohan_verma_LOW_FIT.docx"))
    print("\nAll 3 sample resumes generated successfully.")
    print("Expected outcomes against 'AI/ML Engineer' JD:")
    print("  arjun_sharma_STRONG_FIT.docx  → Strong Fit  (≥75)  → SHORTLIST")
    print("  priya_mehta_AVERAGE_FIT.docx  → Moderate Fit (42-60) → MANUAL REVIEW")
    print("  rohan_verma_LOW_FIT.docx      → Low Fit     (<42)  → REJECT")
